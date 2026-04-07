"""
Extract plain text from DOCX bytes for ingestion.

Design mirrors pdf_extract.py: keep parsing isolated from FastAPI and rag_core so
the pipeline always receives UTF-8 text — the same path as PDF/TXT/MD for chunking,
embeddings, and graph ingest.

Extraction order:
  1. Body paragraphs (preserves document reading order)
  2. Table cells (row-by-row, cells joined with " | ")

Blank paragraphs are skipped; sections are joined with double newlines.
"""
from __future__ import annotations

from io import BytesIO


class DocxExtractError(ValueError):
    """Raised when the file is not a valid DOCX or cannot be read."""


def extract_text_from_docx_bytes(data: bytes) -> str:
    """
    Extract text from a DOCX file given its raw bytes.

    Returns a single string with paragraphs and table rows joined by blank lines.
    Raises DocxExtractError for invalid/empty files.
    """
    if not data:
        raise DocxExtractError("Empty file.")

    try:
        from docx import Document  # python-docx
        doc = Document(BytesIO(data))
    except Exception as e:
        raise DocxExtractError(f"Invalid or corrupted DOCX: {e}") from e

    parts: list[str] = []

    for block in _iter_block_items(doc):
        if isinstance(block, str):
            # Paragraph text
            if block.strip():
                parts.append(block.strip())
        else:
            # Table: emit each row as pipe-separated cells
            for row in block.rows:
                cells = [cell.text.strip() for cell in row.cells]
                # Deduplicate merged cells (python-docx repeats spanned cells)
                unique: list[str] = []
                for c in cells:
                    if not unique or c != unique[-1]:
                        unique.append(c)
                row_text = " | ".join(c for c in unique if c)
                if row_text:
                    parts.append(row_text)

    return "\n\n".join(parts).strip()


def _iter_block_items(doc):
    """
    Yield paragraphs (as str) and tables (as Table objects) in document order.

    python-docx exposes doc.paragraphs and doc.tables separately; to preserve
    reading order we walk the raw XML children of the document body.
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parent_elm = doc.element.body
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent_elm).text
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent_elm)
