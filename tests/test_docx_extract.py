"""Tests for DOCX text extraction (no Vertex / DB)."""
from __future__ import annotations

from io import BytesIO
from unittest.mock import MagicMock, patch

import pytest

from src.docx_extract import DocxExtractError, extract_text_from_docx_bytes


def _make_docx_bytes(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """Build a minimal real DOCX in-memory using python-docx."""
    from docx import Document

    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    if table_rows:
        ncols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=ncols)
        for i, row in enumerate(table_rows):
            for j, cell_text in enumerate(row):
                table.cell(i, j).text = cell_text
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_empty_bytes_raises():
    with pytest.raises(DocxExtractError, match="Empty"):
        extract_text_from_docx_bytes(b"")


def test_extract_invalid_bytes_raises():
    with pytest.raises(DocxExtractError, match="Invalid|corrupted"):
        extract_text_from_docx_bytes(b"not a docx file at all")


def test_extract_paragraphs():
    data = _make_docx_bytes(["Hello world", "Second paragraph"])
    result = extract_text_from_docx_bytes(data)
    assert "Hello world" in result
    assert "Second paragraph" in result


def test_extract_table_rows():
    data = _make_docx_bytes([], table_rows=[["Name", "Age"], ["Alice", "30"], ["Bob", "25"]])
    result = extract_text_from_docx_bytes(data)
    assert "Name" in result
    assert "Alice" in result
    assert "|" in result


def test_extract_paragraphs_and_table_combined():
    data = _make_docx_bytes(
        ["Intro text"],
        table_rows=[["Col A", "Col B"], ["1", "2"]],
    )
    result = extract_text_from_docx_bytes(data)
    assert "Intro text" in result
    assert "Col A" in result


def test_extract_blank_paragraphs_skipped():
    data = _make_docx_bytes(["", "   ", "Real content", ""])
    result = extract_text_from_docx_bytes(data)
    assert "Real content" in result
    # blank lines should not produce leading/trailing whitespace sections
    assert result.strip() == result
